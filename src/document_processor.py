"""
Document processing module for handling various file formats
"""
import os
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
import streamlit as st

# Document processing libraries
try:
    import PyPDF2
    from pdfplumber import PDF
except ImportError:
    PyPDF2 = None
    PDF = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import olefile
except ImportError:
    olefile = None

from utils import clean_text, chunk_text, display_error, display_success, display_info
from config import CHUNK_SIZE, CHUNK_OVERLAP, DOCUMENTS_DIR

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles processing of various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.txt': self._process_txt
        }
    
    def process_file(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """Process a single file and return document metadata and chunks"""
        try:
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                display_error(f"Unsupported file format: {file_extension}")
                return None
            
            display_info(f"Processing {file_path.name}...")
            
            # Extract text content
            text_content = self.supported_formats[file_extension](file_path)
            
            if not text_content:
                display_error(f"No text content extracted from {file_path.name}")
                return None
            
            # Clean and chunk the text
            cleaned_text = clean_text(text_content)
            chunks = chunk_text(cleaned_text, CHUNK_SIZE, CHUNK_OVERLAP)
            
            if not chunks:
                display_error(f"No valid chunks created from {file_path.name}")
                return None
            
            document_data = {
                'file_name': file_path.name,
                'file_path': str(file_path),
                'file_size': file_path.stat().st_size,
                'content': cleaned_text,
                'chunks': chunks,
                'chunk_count': len(chunks),
                'source_type': 'document'
            }
            
            display_success(f"Successfully processed {file_path.name} - {len(chunks)} chunks created")
            return document_data
            
        except Exception as e:
            display_error(f"Error processing {file_path.name}: {str(e)}")
            logger.error(f"Error processing {file_path}: {e}")
            return None
    
    def _process_pdf(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF file"""
        text_content = ""
        
        try:
            # Try with pdfplumber first (better text extraction)
            if PDF:
                with PDF.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
            
            # Fallback to PyPDF2 if pdfplumber fails or isn't available
            elif PyPDF2:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
            else:
                raise ImportError("No PDF processing library available")
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _process_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX file"""
        if not DocxDocument:
            raise ImportError("python-docx library not available")
        
        try:
            doc = DocxDocument(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def _process_doc(self, file_path: Path) -> Optional[str]:
        """Extract text from DOC file (legacy Word format)"""
        # Note: This is a basic implementation. For better DOC support,
        # consider using python-docx2txt or converting to DOCX first
        try:
            # Try to read as plain text (fallback method)
            with open(file_path, 'rb') as file:
                content = file.read()
                # Basic text extraction - this won't work perfectly for all DOC files
                text_content = content.decode('utf-8', errors='ignore')
                # Clean up binary artifacts
                text_content = ''.join(char for char in text_content if char.isprintable() or char.isspace())
                return text_content.strip()
                
        except Exception as e:
            logger.error(f"Error processing DOC {file_path}: {e}")
            # Return empty string rather than failing completely
            display_warning(f"Limited text extraction from {file_path.name}. Consider converting to DOCX for better results.")
            return ""
    
    def _process_txt(self, file_path: Path) -> Optional[str]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error processing TXT {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise
    
    def process_uploaded_file(self, uploaded_file) -> Optional[Dict[str, Any]]:
        """Process a Streamlit uploaded file"""
        try:
            # Save uploaded file temporarily
            temp_path = DOCUMENTS_DIR / uploaded_file.name
            
            with open(temp_path, 'wb') as f:
                f.write(uploaded_file.read())
            
            # Process the file
            result = self.process_file(temp_path)
            
            # Keep the file for future reference
            if result:
                display_success(f"File {uploaded_file.name} saved and processed successfully")
            else:
                # Remove failed file
                temp_path.unlink(missing_ok=True)
            
            return result
            
        except Exception as e:
            display_error(f"Error processing uploaded file {uploaded_file.name}: {str(e)}")
            logger.error(f"Error processing uploaded file: {e}")
            return None
    
    def process_directory(self, directory_path: Path) -> List[Dict[str, Any]]:
        """Process all supported files in a directory"""
        results = []
        
        if not directory_path.exists():
            display_error(f"Directory does not exist: {directory_path}")
            return results
        
        supported_files = []
        for ext in self.supported_formats.keys():
            supported_files.extend(directory_path.glob(f"*{ext}"))
        
        if not supported_files:
            display_info(f"No supported files found in {directory_path}")
            return results
        
        display_info(f"Found {len(supported_files)} files to process")
        
        for file_path in supported_files:
            result = self.process_file(file_path)
            if result:
                results.append(result)
        
        display_success(f"Successfully processed {len(results)} out of {len(supported_files)} files")
        return results
